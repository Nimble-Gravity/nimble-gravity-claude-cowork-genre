# Facilitator tool: regenerate the participant-facing .docx lab files from the
# .txt sources in escape-room/lab-files/. The .txt files are the editable source
# of truth (git-diffable); the .docx files are what participants download —
# Workshop 1/2 feedback: business users recognize Word docs, not .txt/.md.
#
# Usage:  python tools/convert-lab-files.py     (from the escape-room/ folder)
#
# Unlock codes are derived from file CONTENT (e.g. Room 1 counts R. Vega's
# action items), so the conversion is deliberately literal: every non-empty
# line becomes a paragraph, with light heading/bullet styling only.
import os
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
LAB = os.path.join(HERE, "..", "lab-files")

FILES = [
    "vault-standup-notes.txt",
    "raw-notes.txt",
    "interview-notes-jmalik.txt",
    "qa-checklist.txt",
]


def convert(txt_name):
    src = os.path.join(LAB, txt_name)
    out = os.path.join(LAB, txt_name[:-4] + ".docx")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    first_line = True
    with open(src, encoding="utf-8") as f:
        for raw in f.read().splitlines():
            line = raw.rstrip()
            if not line:
                continue
            if line.strip("- ") == "" and "-" in line:  # "---" divider
                doc.add_paragraph("")
                continue
            if first_line:
                doc.add_heading(line, level=1)
                first_line = False
            elif line == line.upper() and any(c.isalpha() for c in line) and len(line) < 60:
                doc.add_heading(line, level=2)
            elif line.startswith("- "):
                doc.add_paragraph(line[2:], style="List Bullet")
            else:
                doc.add_paragraph(line)
    doc.save(out)
    print(f"wrote {os.path.relpath(out, os.path.join(HERE, '..'))}")


for name in FILES:
    convert(name)
